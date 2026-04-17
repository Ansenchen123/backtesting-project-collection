from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"F:\learndata\回測專案整合包")
ASSETS = ROOT / "portfolio_assets"
OUTPUT_DOCX = ROOT / "回測專案作品集.docx"
ACCENT = "102E4A"
ACCENT_LIGHT = "E8F0F7"
TEXT_DARK = RGBColor(32, 39, 49)
TEXT_MUTED = RGBColor(92, 103, 117)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="Microsoft JhengHei", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(paragraph, text, *, size=11, bold=False, color=TEXT_DARK):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return run


def style_paragraph(paragraph, before=0, after=6, line=1.2, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_band(document, title, subtitle):
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(17.4)
    cell = table.cell(0, 0)
    set_cell_shading(cell, ACCENT)
    set_cell_margins(cell, top=140, start=200, bottom=140, end=200)
    p = cell.paragraphs[0]
    add_text(p, title, size=22, bold=True, color=RGBColor(255, 255, 255))
    style_paragraph(p, after=2)
    p2 = cell.add_paragraph()
    add_text(p2, subtitle, size=10, color=RGBColor(222, 233, 242))
    style_paragraph(p2, after=0)
    document.add_paragraph()


def add_info_table(document, rows):
    table = document.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.4)
    table.columns[1].width = Cm(14.0)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], ACCENT_LIGHT)
        set_cell_shading(cells[1], "FFFFFF")
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        p0 = cells[0].paragraphs[0]
        add_text(p0, label, size=10, bold=True, color=RGBColor(22, 62, 99))
        style_paragraph(p0, after=0)
        p1 = cells[1].paragraphs[0]
        add_text(p1, value, size=10)
        style_paragraph(p1, after=0)
    document.add_paragraph()


def add_bullets(document, title, items):
    heading = document.add_paragraph()
    add_text(heading, title, size=13, bold=True, color=RGBColor(22, 62, 99))
    style_paragraph(heading, after=4)
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        add_text(p, item, size=11)
        style_paragraph(p, after=4)


def add_image_with_caption(paragraph_cell, image_path, caption, width=Inches(2.9)):
    if not image_path.exists():
        return
    picture_paragraph = paragraph_cell.paragraphs[0]
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = picture_paragraph.add_run()
    run.add_picture(str(image_path), width=width)
    cap = paragraph_cell.add_paragraph()
    add_text(cap, caption, size=9, color=TEXT_MUTED)
    style_paragraph(cap, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_screenshot_pair(document, left_image, left_caption, right_image, right_caption):
    heading = document.add_paragraph()
    add_text(heading, "操作畫面", size=13, bold=True, color=RGBColor(22, 62, 99))
    style_paragraph(heading, after=4)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.65)
    table.columns[1].width = Cm(8.65)
    left = table.cell(0, 0)
    right = table.cell(0, 1)
    set_cell_margins(left, top=100, start=80, bottom=80, end=80)
    set_cell_margins(right, top=100, start=80, bottom=80, end=80)
    add_image_with_caption(left, left_image, left_caption)
    add_image_with_caption(right, right_image, right_caption)
    document.add_paragraph()


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "第 ", size=9, color=TEXT_MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    set_run_font(run, size=9, color=TEXT_MUTED)
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)
    add_text(p, " 頁", size=9, color=TEXT_MUTED)


def configure_section(section):
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    add_page_number(section)


def add_cover(document):
    add_band(
        document,
        "回測專案作品集",
        "以台股回測、策略模擬與交易驗證腳本為核心的整理版本",
    )

    intro = document.add_paragraph()
    add_text(
        intro,
        "本作品集整理 4 個與回測或交易模擬相關的專案。內容保留各專案原本的策略概念，並補強可執行性、文件完整度與交付結構，讓每個專案都能作為獨立作品進行展示。",
        size=11,
    )
    style_paragraph(intro, after=10)

    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.55)
    table.columns[1].width = Cm(8.55)
    cards = [
        ("作品數量", "4 個獨立專案"),
        ("主題範圍", "追價策略、均線交叉、趨勢回測、Excel 訊號模擬"),
        ("技術基礎", "Python、pandas、matplotlib、FinMind、yfinance"),
        ("整理重點", "中文文件、輸出目錄、可執行入口、過程截圖"),
    ]
    idx = 0
    for r in range(2):
        for c in range(2):
            cell = table.cell(r, c)
            set_cell_shading(cell, "F6F8FB")
            set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
            title, value = cards[idx]
            p1 = cell.paragraphs[0]
            add_text(p1, title, size=10, bold=True, color=RGBColor(22, 62, 99))
            style_paragraph(p1, after=2)
            p2 = cell.add_paragraph()
            add_text(p2, value, size=11)
            style_paragraph(p2, after=0)
            idx += 1
    document.add_paragraph()

    add_bullets(
        document,
        "收錄專案",
        [
            "econamic_model_test：以追價策略為主的台股回測專案",
            "econamic_reload：以均線交叉與帳戶結算邏輯為基礎的模擬交易專案",
            "TB_Strategy：含加碼、停損與最大回撤的趨勢型回測腳本",
            "PY_mod_econamic：以 Excel 訊號資料驅動的交易模擬工具",
        ],
    )


def add_project_page(document, project, first=False):
    if not first:
        section = document.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section)

    add_band(document, project["name"], project["subtitle"])

    heading = document.add_paragraph()
    add_text(heading, "專案簡介", size=13, bold=True, color=RGBColor(22, 62, 99))
    style_paragraph(heading, after=4)
    body = document.add_paragraph()
    add_text(body, project["summary"], size=11)
    style_paragraph(body, after=8)

    add_info_table(
        document,
        [
            ("專案定位", project["positioning"]),
            ("資料來源", project["data_source"]),
            ("主要技術", project["tech"]),
            ("輸出內容", project["outputs"]),
        ],
    )

    add_bullets(document, "整理重點", project["highlights"])
    add_screenshot_pair(
        document,
        project["process_image"],
        project["process_caption"],
        project["result_image"],
        project["result_caption"],
    )


def build():
    document = Document()
    section = document.sections[0]
    configure_section(section)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)

    add_cover(document)
    first_project_section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(first_project_section)

    projects = [
        {
            "name": "econamic_model_test",
            "subtitle": "以追價策略為核心的台股回測專案",
            "summary": "此專案聚焦於一個簡化版追價策略，會抓取歷史價格資料後執行回測，輸出交易紀錄、資金曲線與圖表，適合作為策略原型驗證與回測流程展示。",
            "positioning": "策略原型驗證 / 台股回測腳本",
            "data_source": "tvDatafeed，並補上 yfinance fallback",
            "tech": "Python、pandas、matplotlib、yfinance、tvDatafeed",
            "outputs": "歷史資料 CSV、交易紀錄 CSV、資金曲線 CSV、回測圖表",
            "highlights": [
                "將原本較實驗性的策略腳本整理成可交付的命令列專案。",
                "保留追價策略概念，但補上輸出目錄、文件與執行入口。",
                "適合展示從資料下載到回測輸出的完整流程。",
            ],
            "process_image": ASSETS / "econamic_model_test_process.png",
            "process_caption": "執行過程截圖：完成回測後的命令列輸出",
            "result_image": ROOT / "econamic_model_test" / "outputs" / "2330_chart.png",
            "result_caption": "結果截圖：收盤價與策略權益曲線圖",
        },
        {
            "name": "econamic_reload",
            "subtitle": "以均線交叉與帳戶結算邏輯為基礎的模擬交易專案",
            "summary": "此專案以短長均線交叉作為進出場條件，並保留帳戶、交割延遲與每日結算概念，不只計算最終報酬，也模擬了更接近實際帳戶流轉的過程。",
            "positioning": "模擬交易流程 / 台股均線回測",
            "data_source": "FinMind，含本地快取",
            "tech": "Python、pandas、matplotlib、FinMind",
            "outputs": "資金曲線 CSV、帳戶摘要 CSV、帳戶 JSON、回測圖表、資料快取",
            "highlights": [
                "把資料下載、帳戶管理與回測輸出整理到同一套可執行流程中。",
                "保留每日結算概念，使專案更像模擬交易而非單純績效計算。",
                "適合作為較完整的回測小工具展示。",
            ],
            "process_image": ASSETS / "econamic_reload_process.png",
            "process_caption": "執行過程截圖：模擬完成後的摘要輸出",
            "result_image": ROOT / "econamic_reload" / "outputs" / "2330_backtest.png",
            "result_caption": "結果截圖：價格、均線與資產曲線圖",
        },
        {
            "name": "TB_Strategy",
            "subtitle": "含加碼、停損與最大回撤的趨勢型回測腳本",
            "summary": "此專案以 20MA 與 60MA 作為趨勢判斷基礎，並加入加碼、停損與最大回撤計算。整理後保留了策略本體的特性，同時補上參數入口與結果輸出，讓它更適合作為展示用腳本。",
            "positioning": "趨勢追蹤回測 / 參數化策略腳本",
            "data_source": "tvDatafeed，並補上 yfinance fallback",
            "tech": "Python、pandas、matplotlib、yfinance、tvDatafeed",
            "outputs": "交易紀錄 CSV、績效摘要 CSV、權益曲線圖",
            "highlights": [
                "保留加碼與停損邏輯，使策略不只停留在單純均線交叉。",
                "補上命令列參數與固定輸出路徑，降低原本腳本化的使用門檻。",
                "適合作為較偏策略研究風格的作品。",
            ],
            "process_image": ASSETS / "TB_Strategy_process.png",
            "process_caption": "執行過程截圖：回測完成後的績效摘要",
            "result_image": ROOT / "TB_Strategy" / "outputs" / "2330_equity_curve.png",
            "result_caption": "結果截圖：權益曲線與峰值比較圖",
        },
        {
            "name": "PY_mod_econamic",
            "subtitle": "以 Excel 訊號資料驅動的交易模擬工具",
            "summary": "此專案不是從零建立訊號，而是讀取外部整理好的價格與訊號資料，進行模擬交易與資產曲線比較。它更偏向策略訊號驗證工具，適合用於測試既有分類器或判斷模型的輸出效果。",
            "positioning": "訊號驗證工具 / Excel 驅動交易模擬",
            "data_source": "本地 Excel 價格資料與訊號資料",
            "tech": "Python、pandas、matplotlib、openpyxl",
            "outputs": "每日模擬結果 CSV、摘要 CSV、模擬圖表",
            "highlights": [
                "強調將外部訊號資料快速轉成策略模擬結果的能力。",
                "可直接比較策略表現與 Buy and Hold，方便觀察訊號有效性。",
                "在整理後更適合作為驗證型工具展示。",
            ],
            "process_image": ASSETS / "PY_mod_econamic_process.png",
            "process_caption": "執行過程截圖：模擬完成後的命令列結果",
            "result_image": ROOT / "PY_mod_econamic" / "outputs" / "google_simulation.png",
            "result_caption": "結果截圖：策略資產曲線與 Buy and Hold 比較圖",
        },
    ]

    for idx, project in enumerate(projects):
        add_project_page(document, project, first=(idx == 0))

    document.save(str(OUTPUT_DOCX))


if __name__ == "__main__":
    build()
